import docx

# 문서 생성
doc = docx.Document()

# 단락 추가
par = doc.add_paragraph('par 단락을 추가합니다.')
par2 = par.insert_paragraph_before('par 단락 앞에 새 단락을 추가합니다.')

# 테이블 추가
table = doc.add_table(rows=1, cols=3, style="Table Grid")
table.cell(0, 0).text = "(0,0)셀"
table.rows[0].cells[1].text = "(0,1)셀"
table.columns[2].cells[0].text = "(0,2)셀"

# 제목 추가
for i in range(10):
    doc.add_heading('level{} 제목을 추가합니다.'.format(i), level=i)
# 문서 저장
doc.save('output\\docx_write.docx')
