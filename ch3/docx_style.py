import docx

# 문서 생성
doc = docx.Document()
# 단락 추가
par = doc.add_paragraph()
# Run 객체 추가
for i in range(5):
    par.add_run('{})par 단락에 {}번째 run을 추가했습니다.'.format(i + 1, i + 1))
    r_num = len(par.runs)
    print('({})run 갯수: {}'.format(i + 1, r_num))
# 각 Run 객체에 font 속성 설정
r = par.runs
r[0].font.bold = True
r[1].font.italic = True
r[2].font.underline = True
r[3].font.size = docx.shared.Pt(20)
r[4].font.color.rgb = docx.shared.RGBColor(255, 0, 0)

# 문서 저장
doc.save('./output\\docx_stly.docx')
