import docx

# 기존 원드 파일 열기
doc = docx.Document('./output/docx_write.docx')

# 단락 수정
doc.paragraphs[0].text = '첫 번째 단락을 교체합니다.'

# 테이블 수정
table = doc.tables[0]
new_row = table.add_row()
r_num = len(table.rows)
c_num = len(table.columns)
for c in range(c_num):
    new_row.cells[c].text = "({},{})셀 추가하기".format(r_num, c)
# 단락과 테이블 출력
for p in doc.paragraphs:
    print(p.text)
for r in range(r_num):
    for c in range(c_num):
        print(table.cell(r, c).text)
doc.save('./output/docx_read.docx')
